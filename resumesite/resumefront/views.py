from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.decorators import login_required
from django.contrib.auth.mixins import LoginRequiredMixin
from django.contrib import messages
from django.views.generic import ListView, DetailView, CreateView, UpdateView, DeleteView
from django.urls import reverse_lazy
from django.http import JsonResponse
from docx import Document
from docx.shared import Pt
from io import BytesIO
import logging
import requests
import os
from dotenv import load_dotenv

load_dotenv()


from django.http import HttpResponse


logging.basicConfig(level=logging.DEBUG)
logger = logging.getLogger(__name__)





from .models import (
    UserProfile,
    Resume,
    Education,
    Experience,
    Skill,
    Project,
    Certificate,
    Language
)
from .forms import (
    UserRegistrationForm,
    UserProfileForm,
    ResumeForm,
    EducationForm,
    ExperienceForm,
    SkillForm,
    ProjectForm,
    CertificateForm,
    LanguageForm,
    ResumeImportForm,
    AIGenerateForm
)

def index(request):
    """Home page view"""
    return render(request, 'index.html')

def register(request):
    """User registration view"""
    if request.method == 'POST':
        user_form = UserRegistrationForm(request.POST)
        profile_form = UserProfileForm(request.POST, request.FILES)
        
        if user_form.is_valid() and profile_form.is_valid():
            user = user_form.save()
            profile = profile_form.save(commit=False)
            profile.user = user
            profile.save()
            
            messages.success(request, 'Registration successful! Please log in.')
            return redirect('login')
    else:
        user_form = UserRegistrationForm()
        profile_form = UserProfileForm()
    
    return render(request, 'registration/register.html', {
        'user_form': user_form,
        'profile_form': profile_form
    })

@login_required
def dashboard(request):
    """User dashboard showing all resumes"""
    resumes = Resume.objects.filter(user=request.user)
    return render(request, 'dashboard/dashboard.html', {'resumes': resumes})

class ResumeCreateView(LoginRequiredMixin, CreateView):
    """Create a new resume"""
    model = Resume
    form_class = ResumeForm
    template_name = 'resume/create.html'
    success_url = reverse_lazy('dashboard')

    def form_valid(self, form):
        form.instance.user = self.request.user
        return super().form_valid(form)

class ResumeDetailView(LoginRequiredMixin, DetailView):
    """View a complete resume"""
    model = Resume
    template_name = 'resume/detail.html'
    context_object_name = 'resume'

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        resume = self.get_object()
        context.update({
            'education_list': resume.education.all(),
            'experience_list': resume.experiences.all(),
            'skill_list': resume.skills.all(),
            'project_list': resume.projects.all(),
            'certificate_list': resume.certificates.all(),
            'language_list': resume.languages.all(),
        })
        return context

@login_required
def resume_edit(request, pk):
    """Edit resume sections"""
    resume = get_object_or_404(Resume, pk=pk, user=request.user)
    
    if request.method == 'POST':
        form = ResumeForm(request.POST, instance=resume)
        if form.is_valid():
            form.save()
            messages.success(request, 'Resume updated successfully!')
            return redirect('resume_detail', pk=pk)
    else:
        form = ResumeForm(instance=resume)
    
    return render(request, 'resume/edit.html', {
        'form': form,
        'resume': resume
    })

@login_required
def section_add(request, pk, section_type):
    """Add a new section to resume"""
    resume = get_object_or_404(Resume, pk=pk, user=request.user)
    
    form_classes = {
        'education': EducationForm,
        'experience': ExperienceForm,
        'skill': SkillForm,
        'project': ProjectForm,
        'certificate': CertificateForm,
        'language': LanguageForm
    }
    
    FormClass = form_classes.get(section_type)
    if not FormClass:
        messages.error(request, 'Invalid section type')
        return redirect('resume_detail', pk=pk)
    
    if request.method == 'POST':
        form = FormClass(request.POST)
        if form.is_valid():
            section = form.save(commit=False)
            section.resume = resume
            section.save()
            messages.success(request, f'{section_type.title()} added successfully!')
            return redirect('resume_detail', pk=pk)
    else:
        form = FormClass()
    
    return render(request, 'resume/section_add.html', {
        'form': form,
        'resume': resume,
        'section_type': section_type
    })

@login_required
def section_edit(request, pk, section_type, section_id):
    """Edit a specific section of resume"""
    resume = get_object_or_404(Resume, pk=pk, user=request.user)
    
    model_classes = {
        'education': Education,
        'experience': Experience,
        'skill': Skill,
        'project': Project,
        'certificate': Certificate,
        'language': Language
    }
    
    form_classes = {
        'education': EducationForm,
        'experience': ExperienceForm,
        'skill': SkillForm,
        'project': ProjectForm,
        'certificate': CertificateForm,
        'language': LanguageForm
    }
    
    ModelClass = model_classes.get(section_type)
    FormClass = form_classes.get(section_type)
    
    if not ModelClass or not FormClass:
        messages.error(request, 'Invalid section type')
        return redirect('resume_detail', pk=pk)
    
    section = get_object_or_404(ModelClass, pk=section_id, resume=resume)
    
    if request.method == 'POST':
        form = FormClass(request.POST, instance=section)
        if form.is_valid():
            form.save()
            messages.success(request, f'{section_type.title()} updated successfully!')
            return redirect('resume_detail', pk=pk)
    else:
        form = FormClass(instance=section)
    
    return render(request, 'resume/section_edit.html', {
        'form': form,
        'resume': resume,
        'section_type': section_type,
        'section': section
    })

@login_required
def section_delete(request, pk, section_type, section_id):
    """Delete a specific section from resume"""
    resume = get_object_or_404(Resume, pk=pk, user=request.user)
    
    model_classes = {
        'education': Education,
        'experience': Experience,
        'skill': Skill,
        'project': Project,
        'certificate': Certificate,
        'language': Language
    }
    
    ModelClass = model_classes.get(section_type)
    if not ModelClass:
        messages.error(request, 'Invalid section type')
        return redirect('resume_detail', pk=pk)
    
    section = get_object_or_404(ModelClass, pk=section_id, resume=resume)
    
    if request.method == 'POST':
        section.delete()
        messages.success(request, f'{section_type.title()} deleted successfully!')
        return redirect('resume_detail', pk=pk)
    
    return render(request, 'resume/section_delete.html', {
        'resume': resume,
        'section_type': section_type,
        'section': section
    })

@login_required
def resume_import(request):
    """Import resume from file"""
    if request.method == 'POST':
        form = ResumeImportForm(request.POST, request.FILES)
        if form.is_valid():
            # Here you would add your AI processing logic
            messages.success(request, 'Resume imported successfully!')
            return redirect('dashboard')
    else:
        form = ResumeImportForm()
    
    return render(request, 'resumefront/templates/resume/import.html', {'form': form})



@login_required
def enhance_summary(request, resume_id):
    logger.debug(f"Enhance summary called for resume ID: {resume_id}")
    resume = get_object_or_404(Resume, id=resume_id, user=request.user)
    
    if request.method == "POST":
        logger.debug("Processing POST request...")
        
        # Check if this is a download request
        is_download = 'download' in request.POST
        
        # Fetch all resume categories
        raw_summary = resume.summary or "No summary provided."
        experience_list = resume.experiences.all()
        education_list = resume.education.all()
        skill_list = resume.skills.all()
        project_list = resume.projects.all()
        
        # If it's a download request and we already have a summary, skip enhancement
        if is_download and raw_summary != "No summary provided.":
            logger.debug("Download request with existing summary, skipping enhancement")
            return generate_word_document(resume, experience_list, education_list, skill_list, project_list)
        
        # Prepare context for AI enhancement
        # resume_context = {
        #     "experiences": [{"position": exp.position, "company": exp.company, 
        #                     "start_date": exp.start_date, "end_date": "Present" if exp.is_current else exp.end_date,
        #                     "description": exp.description} for exp in experience_list],
            # "education": [{"degree": edu.degree, "field": edu.field_of_study, 
            #               "institution": edu.institution} for edu in education_list],
            # "skills": [{"name": skill.name, "level": skill.get_level_display()} for skill in skill_list]
        # }
        # Prepare context for AI enhancement
        resume_context = {
            "experiences": [{"position": exp.position, "company": exp.company, 
                            "start_date": exp.start_date, "end_date": exp.end_date,
                            "is_current": exp.is_current, # Make sure this property exists
                            "description": exp.description} for exp in experience_list],
            "education": [{"degree": edu.degree, "field": edu.field_of_study, 
                          "institution": edu.institution} for edu in education_list],
            "skills": [{"name": skill.name, "level": skill.get_level_display()} for skill in skill_list]
        }
        
        # Call AI API to improve the summary
        try:
            polished_summary = call_gemini_api(raw_summary, resume_context)
            logger.debug(f"AI-enhanced summary: {polished_summary}")
            
            # Save the enhanced summary back to the resume
            resume.summary = polished_summary
            resume.save()
            
        except Exception as e:
            error_message = str(e)
            logger.error(f"Error calling AI API: {error_message}")
            
            messages.error(request, f"Error enhancing summary: {error_message}")
            return redirect('resume_detail', pk=resume_id)
        
        # If this is a download request, generate and return the Word document
        if is_download:
            return generate_word_document(resume, experience_list, education_list, skill_list, project_list)
        
        # If no download requested, redirect back to the resume detail page with success message
        messages.success(request, "Resume summary successfully enhanced!")
        return redirect('resume_detail', pk=resume_id)

    # If GET request, just render the template
    return render(request, 'resume/detail.html', {'resume': resume})

def generate_word_document(resume, experience_list, education_list, skill_list, project_list):
    """Generate and return a Word document with the resume content"""
    # Generate Word Document
    doc = Document()
    
    # Add a title
    doc.add_heading(resume.title, 0)

    # Add Professional Summary section
    doc.add_heading("Professional Summary", level=1)
    p = doc.add_paragraph(resume.summary)
    p.style.font.size = Pt(12)
    
    # Add Experience section if available
    if experience_list:
        doc.add_heading("Experience", level=1)
        for exp in experience_list:
            p = doc.add_paragraph(
                f"- {exp.position} at {exp.company} ({exp.start_date} - "
                f"{'Present' if exp.is_current else exp.end_date}): {exp.description}",
                style='List Bullet'
            )
            p.style.font.size = Pt(11)
    
    # Add Education section if available
    if education_list:
        doc.add_heading("Education", level=1)
        for edu in education_list:
            p = doc.add_paragraph(
                f"- {edu.degree} in {edu.field_of_study}, {edu.institution} "
                f"({edu.start_date} - {edu.end_date}): {edu.description or 'N/A'}",
                style='List Bullet'
            )
            p.style.font.size = Pt(11)
    
    # Add Skills section if available
    if skill_list:
        doc.add_heading("Skills", level=1)
        for skill in skill_list:
            p = doc.add_paragraph(
                f"- {skill.name} ({skill.get_level_display()})",
                style='List Bullet'
            )
            p.style.font.size = Pt(11)
    
    # Add Projects section if available
    if project_list:
        doc.add_heading("Projects", level=1)
        for project in project_list:
            p = doc.add_paragraph(
                f"- {project.name}: {project.description}",
                style='List Bullet'
            )
            p.style.font.size = Pt(11)
    
    # Save the document to a BytesIO object
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    # Create a response to download the Word file
    response = HttpResponse(
        buffer.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = f'attachment; filename="Polished_Resume_{resume.title}.docx"'
    return response




def call_gemini_api(summary, resume_context):
    """
    Use Google's Gemini API to enhance resume summary using Gemini 2.0 Flash model
    """
    import google.generativeai as genai
    import random
    
    # Set up the API key
    GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
    if not GEMINI_API_KEY:
        logger.error("GEMINI_API_KEY environment variable is not set")
        return generate_fallback_summary(resume_context)
        
    # Configure the API
    genai.configure(api_key=GEMINI_API_KEY)
    
    # Log API key status
    logger.debug(f"Gemini API key configured: {'Yes' if GEMINI_API_KEY else 'No'}")
    
    # Format resume context
    experience_text = "\n".join([
        f"- {exp['position']} at {exp['company']} ({exp['start_date']} - "
        f"{'Present' if exp.get('is_current', False) else exp.get('end_date', 'Not specified')})\n"
        f"  Description: {exp['description']}" 
        for exp in resume_context['experiences'][:3]
    ])
    
    skills_text = ", ".join([skill['name'] for skill in resume_context['skills'][:7]])
    
    education_text = "\n".join([
        f"- {edu['degree']} in {edu['field']} from {edu['institution']}" 
        for edu in resume_context['education'][:2]
    ])
    
    # Create unique identifier
    unique_id = random.randint(1000, 9999)
    
    # Create prompt
    prompt = f"""
Generate a professional resume summary for someone with the following background:

Experience:
{experience_text}

Skills: {skills_text}

Education:
{education_text}

Current summary: {summary}

IMPORTANT: Respond ONLY with the finished summary in 3-4 sentences. DO NOT include explanations, options, or meta-commentary. DO NOT refer to the resume by number or ID. DO NOT ask questions. Just write a polished, professional summary that would appear at the top of a resume.
"""
    
    try:
        # Use Gemini 2.0 Flash model (based on your API key setup)
        logger.info("Using gemini-2.0-flash model")
        model = genai.GenerativeModel("gemini-2.0-flash")
        
        try:
            # Try with basic configuration
            generation_config = {
                "temperature": 0.9,
                "top_p": 0.95,
                "max_output_tokens": 250,
            }
            
            response = model.generate_content(prompt, generation_config=generation_config)
            enhanced_summary = response.text.strip()
            
            # If we get a good response, return it
            if enhanced_summary:
                logger.info(f"Generated summary: {enhanced_summary[:100]}...")
                return enhanced_summary
                
        except Exception as first_error:
            logger.warning(f"First attempt failed: {str(first_error)}")
            
            # Try with simpler configuration
            try:
                response = model.generate_content(prompt)
                enhanced_summary = response.text.strip()
                
                if enhanced_summary:
                    logger.info(f"Generated summary (second attempt): {enhanced_summary[:100]}...")
                    return enhanced_summary
                    
            except Exception as second_error:
                logger.warning(f"Second attempt failed: {str(second_error)}")
                
                # One last attempt with direct API call if library is failing
                try:
                    import requests
                    import json
                    
                    api_url = "https://generativelanguage.googleapis.com/v1beta/models/gemini-2.0-flash:generateContent"
                    headers = {
                        "Content-Type": "application/json"
                    }
                    
                    data = {
                        "contents": [
                            {
                                "parts": [
                                    {
                                        "text": prompt
                                    }
                                ]
                            }
                        ]
                    }
                    
                    params = {
                        "key": GEMINI_API_KEY
                    }
                    
                    response = requests.post(api_url, headers=headers, json=data, params=params)
                    
                    if response.status_code == 200:
                        result = response.json()
                        if "candidates" in result and len(result["candidates"]) > 0:
                            if "content" in result["candidates"][0] and "parts" in result["candidates"][0]["content"]:
                                parts = result["candidates"][0]["content"]["parts"]
                                if len(parts) > 0 and "text" in parts[0]:
                                    enhanced_summary = parts[0]["text"].strip()
                                    logger.info(f"Generated summary (direct API): {enhanced_summary[:100]}...")
                                    return enhanced_summary
                    
                    logger.warning(f"Direct API call failed with status: {response.status_code}")
                    
                except Exception as third_error:
                    logger.error(f"All generation attempts failed: {str(third_error)}")
        
        # If we reach here, all attempts failed
        logger.warning("All generation attempts failed, using fallback")
        return generate_fallback_summary(resume_context)
            
    except Exception as e:
        logger.error(f"Error setting up Gemini API: {str(e)}")
        return generate_fallback_summary(resume_context)




def generate_fallback_summary(resume_context):
    """Generate a fallback summary when API calls fail"""
    # Create a basic summary based on the available information
    skills_part = ""
    if resume_context['skills']:
        top_skills = [skill['name'] for skill in resume_context['skills'][:3]]
        skills_part = f"Skilled in {', '.join(top_skills)}"
    
    experience_part = ""
    if resume_context['experiences']:
        exp = resume_context['experiences'][0]
        experience_part = f" with experience as {exp['position']} at {exp['company']}"
    
    education_part = ""
    if resume_context['education']:
        edu = resume_context['education'][0]
        education_part = f" Holds a {edu['degree']} in {edu['field']} from {edu['institution']}."
    
    # Assemble the summary
    summary = f"Dedicated professional {experience_part}. {skills_part}. {education_part} Committed to delivering high-quality results with a proven track record of success."
    
    return summary


@login_required
def ai_generate(request):
    """Generate content using AI"""
    if request.method == 'POST':
        form = AIGenerateForm(request.POST)
        if form.is_valid():
            # Here you would add your AI generation logic
            return JsonResponse({'success': True, 'generated_content': 'AI generated content would go here'})
    else:
        form = AIGenerateForm()
    
    return render(request, 'resumefront/templates/resume/ai_generate.html', {'form': form})

class ResumeDeleteView(LoginRequiredMixin, DeleteView):
    """Delete a resume"""
    model = Resume
    template_name = 'resume/delete.html'
    success_url = reverse_lazy('dashboard')

    def get_queryset(self):
        return Resume.objects.filter(user=self.request.user)